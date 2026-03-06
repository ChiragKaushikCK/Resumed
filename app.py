import streamlit as st
import os
import json
import io
import re
import pandas as pd
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# ==========================================
# 1. API Configuration (OpenRouter)
# ==========================================
try:
    api_key = st.secrets["OPENROUTER_API_KEY"]
except (FileNotFoundError, KeyError):
    api_key = os.getenv("OPENROUTER_API_KEY")

if not api_key:
    st.error("⚠️ API key not found. Please set it in Streamlit Secrets or as an environment variable.")
    st.stop()

client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=api_key,
)

# ==========================================
# 2. HTML/CSS Resume Templates (Preview Only)
# ==========================================
def render_template(data, style="faang"):
    accent = "#1a1a1a" if style == "faang" else "#1d5fa8"
    font = "'Helvetica Neue', Helvetica, Arial, sans-serif" if style == "faang" else "'Georgia', serif"

    def section_header(title):
        return f"""<h3 style="border-bottom: 2px solid {accent}; padding-bottom: 3px; margin-top: 20px;
                   text-transform: uppercase; font-size: 13px; color: {accent}; letter-spacing: 1px;">{title}</h3>"""

    html = f"""
    <div style="font-family: {font}; max-width: 780px; margin: 0 auto; padding: 36px;
                color: #1a1a1a; background: white; box-shadow: 0 2px 16px rgba(0,0,0,0.10);">
        <div style="text-align: center; padding-bottom: 12px; margin-bottom: 4px; border-bottom: 2px solid {accent};">
            <h1 style="margin: 0; font-size: 30px; letter-spacing: 0.5px;">{data.get('name', 'Your Name')}</h1>
            <p style="margin: 6px 0 0 0; font-size: 12px; color: #444;">{data.get('contact', '')}</p>
        </div>
    """

    if data.get('summary'):
        html += section_header("Professional Summary")
        html += f'<p style="font-size: 12.5px; line-height: 1.6; margin-top: 6px;">{data["summary"]}</p>'

    exp_list = [e for e in data.get('experience', []) if e.get('title') or e.get('company') or e.get('description')]
    if exp_list:
        html += section_header("Experience")
        for exp in exp_list:
            desc_html = exp.get('description', '').replace('\n', '<br>')
            # Convert bullet lines to HTML bullets
            desc_html = re.sub(r'(?m)^[•\-\*]\s?', '• ', desc_html)
            html += f"""
            <div style="margin-top: 10px;">
                <div style="display: flex; justify-content: space-between; align-items: baseline;">
                    <span><b style="font-size: 13px;">{exp.get('title', '')} — {exp.get('company', '')}</b></span>
                    <span style="font-size: 11.5px; color: #666; white-space: nowrap; margin-left: 12px;">{exp.get('duration', '')}</span>
                </div>
                <p style="font-size: 12px; margin: 4px 0; line-height: 1.6; color: #333;">{desc_html}</p>
            </div>"""

    proj_list = [p for p in data.get('projects', []) if p.get('name') or p.get('description')]
    if proj_list:
        html += section_header("Projects")
        for proj in proj_list:
            html += f"""
            <div style="margin-top: 10px;">
                <div style="display: flex; justify-content: space-between; align-items: baseline;">
                    <b style="font-size: 13px;">{proj.get('name', '')}</b>
                    <span style="font-size: 11.5px; color: #666; font-style: italic; margin-left: 12px;">{proj.get('tech_stack', '')}</span>
                </div>
                <p style="font-size: 12px; margin: 4px 0; line-height: 1.6; color: #333;">{proj.get('description', '')}</p>
            </div>"""

    edu_list = [e for e in data.get('education', []) if e.get('degree') or e.get('university')]
    if edu_list:
        html += section_header("Education")
        for edu in edu_list:
            html += f"""
            <div style="display: flex; justify-content: space-between; margin-top: 8px; align-items: baseline;">
                <div>
                    <b style="font-size: 13px;">{edu.get('university', '')}</b>
                    <p style="margin: 1px 0; font-size: 12px; color: #444;">{edu.get('degree', '')}</p>
                </div>
                <span style="font-size: 11.5px; color: #666; white-space: nowrap; margin-left: 12px;">{edu.get('year', '')}</span>
            </div>"""

    if data.get('skills', '').strip():
        html += section_header("Skills")
        html += f'<p style="font-size: 12.5px; line-height: 1.6; margin-top: 6px;">{data["skills"]}</p>'

    html += "</div>"
    return html

# ==========================================
# 3. AI Processing
# ==========================================
def extract_details_with_ai(raw_text):
    prompt = """
You are an expert resume writer. Extract information from the user's raw text and return ONLY a valid JSON object (no markdown, no code fences).

RULES:
1. Only include sections the user actually mentioned. If no experience is provided, set "experience" to an empty array [].
2. If experience is mentioned but sparse, expand into professional bullet points.
3. Only include "projects" if the user mentioned projects. Otherwise set to [].
4. Write a compelling summary if one isn't provided.
5. Infer skills from context if not explicitly listed.
6. Descriptions should use bullet-point style (start lines with •).

JSON Schema (return exactly this structure):
{
  "name": "Full Name",
  "contact": "Email | Phone | Location | LinkedIn",
  "summary": "2-3 sentence professional summary.",
  "experience": [
    {
      "title": "Job Title",
      "company": "Company",
      "duration": "Month Year – Month Year",
      "description": "• Achievement 1\n• Achievement 2\n• Achievement 3"
    }
  ],
  "projects": [
    {
      "name": "Project Name",
      "tech_stack": "Tech 1, Tech 2",
      "description": "What it does and its impact."
    }
  ],
  "education": [
    {
      "degree": "Degree, Major",
      "university": "University Name",
      "year": "Graduation Year"
    }
  ],
  "skills": "Skill 1, Skill 2, Skill 3, ..."
}
"""
    try:
        response = client.chat.completions.create(
            model="openai/gpt-4o-mini",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": raw_text}
            ],
            response_format={"type": "json_object"}
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        st.error(f"Error communicating with OpenRouter API: {e}")
        return None

# ==========================================
# 4. High-Quality PDF via ReportLab
# ==========================================
def generate_pdf_reportlab(data, style="faang"):
    accent_color = colors.HexColor("#1a1a1a") if style == "faang" else colors.HexColor("#1d5fa8")
    font_name = "Helvetica" if style == "faang" else "Times-Roman"
    font_bold = "Helvetica-Bold" if style == "faang" else "Times-Bold"
    font_italic = "Helvetica-Oblique" if style == "faang" else "Times-Italic"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18*mm, rightMargin=18*mm,
        topMargin=16*mm, bottomMargin=16*mm
    )

    styles = {
        "name": ParagraphStyle("name", fontName=font_bold, fontSize=22, alignment=TA_CENTER,
                               textColor=colors.HexColor("#1a1a1a"), spaceAfter=3),
        "contact": ParagraphStyle("contact", fontName=font_name, fontSize=9.5, alignment=TA_CENTER,
                                  textColor=colors.HexColor("#555555"), spaceAfter=8),
        "section": ParagraphStyle("section", fontName=font_bold, fontSize=10, textColor=accent_color,
                                  spaceBefore=12, spaceAfter=2, letterSpacing=1.5),
        "body": ParagraphStyle("body", fontName=font_name, fontSize=10, leading=15,
                               textColor=colors.HexColor("#222222"), spaceAfter=2),
        "job_title": ParagraphStyle("job_title", fontName=font_bold, fontSize=10.5,
                                    textColor=colors.HexColor("#1a1a1a"), spaceAfter=1),
        "bullet": ParagraphStyle("bullet", fontName=font_name, fontSize=10, leading=14,
                                 textColor=colors.HexColor("#333333"), leftIndent=12, spaceAfter=1),
        "italic": ParagraphStyle("italic", fontName=font_italic, fontSize=9.5,
                                 textColor=colors.HexColor("#666666"), spaceAfter=4),
    }

    def hr():
        return HRFlowable(width="100%", thickness=1, color=accent_color, spaceAfter=4, spaceBefore=2)

    def section_title(title):
        return [Paragraph(title.upper(), styles["section"]), hr()]

    story = []

    # Header
    story.append(Paragraph(data.get("name", "Your Name"), styles["name"]))
    story.append(Paragraph(data.get("contact", ""), styles["contact"]))
    story.append(HRFlowable(width="100%", thickness=2, color=accent_color, spaceAfter=6))

    # Summary
    if data.get("summary", "").strip():
        story += section_title("Professional Summary")
        story.append(Paragraph(data["summary"], styles["body"]))

    # Experience
    exp_list = [e for e in data.get("experience", []) if e.get("title") or e.get("company") or e.get("description")]
    if exp_list:
        story += section_title("Experience")
        for exp in exp_list:
            # Title + duration on same row
            title_str = f"{exp.get('title', '')} — <b>{exp.get('company', '')}</b>"
            dur_str = exp.get("duration", "")
            row = Table(
                [[Paragraph(title_str, styles["job_title"]), Paragraph(dur_str, styles["italic"])]],
                colWidths=[None, 45*mm]
            )
            row.setStyle(TableStyle([
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(row)
            for line in exp.get("description", "").split("\n"):
                line = line.strip()
                if line:
                    clean = re.sub(r'^[•\-\*]\s?', '', line)
                    story.append(Paragraph(f"• {clean}", styles["bullet"]))
            story.append(Spacer(1, 4))

    # Projects
    proj_list = [p for p in data.get("projects", []) if p.get("name") or p.get("description")]
    if proj_list:
        story += section_title("Projects")
        for proj in proj_list:
            name_str = f"<b>{proj.get('name', '')}</b>"
            tech_str = proj.get("tech_stack", "")
            row = Table(
                [[Paragraph(name_str, styles["job_title"]), Paragraph(tech_str, styles["italic"])]],
                colWidths=[None, 60*mm]
            )
            row.setStyle(TableStyle([
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(row)
            story.append(Paragraph(proj.get("description", ""), styles["body"]))
            story.append(Spacer(1, 4))

    # Education
    edu_list = [e for e in data.get("education", []) if e.get("degree") or e.get("university")]
    if edu_list:
        story += section_title("Education")
        for edu in edu_list:
            uni_str = f"<b>{edu.get('university', '')}</b>"
            year_str = edu.get("year", "")
            row = Table(
                [[Paragraph(uni_str, styles["job_title"]), Paragraph(year_str, styles["italic"])]],
                colWidths=[None, 35*mm]
            )
            row.setStyle(TableStyle([
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(row)
            story.append(Paragraph(edu.get("degree", ""), styles["body"]))
            story.append(Spacer(1, 4))

    # Skills
    if data.get("skills", "").strip():
        story += section_title("Skills")
        story.append(Paragraph(data["skills"], styles["body"]))

    doc.build(story)
    return buf.getvalue()

# ==========================================
# 5. High-Quality DOCX Generator
# ==========================================
def add_horizontal_rule(doc, color_hex="1a1a1a", thickness=4):
    """Adds a styled bottom border to the last paragraph as a divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def set_font(run, name, size, bold=False, italic=False, color_hex=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)

def generate_docx(data, style="faang"):
    accent_hex = "1a1a1a" if style == "faang" else "1d5fa8"
    main_font = "Calibri"

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data.get("name", "Your Name"))
    set_font(name_run, main_font, 24, bold=True, color_hex="1a1a1a")
    name_para.paragraph_format.space_after = Pt(2)

    # Contact
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = contact_para.add_run(data.get("contact", ""))
    set_font(cr, main_font, 9.5, color_hex="555555")
    contact_para.paragraph_format.space_after = Pt(6)
    add_horizontal_rule(doc, accent_hex, thickness=12)

    def section_heading(title):
        p = doc.add_paragraph()
        r = p.add_run(title.upper())
        set_font(r, main_font, 10, bold=True, color_hex=accent_hex)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        add_horizontal_rule(doc, accent_hex, thickness=4)

    def body_para(text, indent=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, main_font, 10, color_hex="222222")
        p.paragraph_format.space_after = Pt(1)
        if indent:
            p.paragraph_format.left_indent = Cm(0.4)
        return p

    # Summary
    if data.get("summary", "").strip():
        section_heading("Professional Summary")
        body_para(data["summary"])

    # Experience
    exp_list = [e for e in data.get("experience", []) if e.get("title") or e.get("company") or e.get("description")]
    if exp_list:
        section_heading("Experience")
        for exp in exp_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            r1 = p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            set_font(r1, main_font, 11, bold=True, color_hex="1a1a1a")
            r2 = p.add_run(f"   {exp.get('duration', '')}")
            set_font(r2, main_font, 9.5, italic=True, color_hex="666666")
            for line in exp.get("description", "").split("\n"):
                line = line.strip()
                if line:
                    clean = re.sub(r'^[•\-\*]\s?', '', line)
                    bp = doc.add_paragraph(style="List Bullet")
                    r = bp.add_run(clean)
                    set_font(r, main_font, 10, color_hex="333333")
                    bp.paragraph_format.left_indent = Cm(0.5)
                    bp.paragraph_format.space_after = Pt(1)

    # Projects
    proj_list = [p for p in data.get("projects", []) if p.get("name") or p.get("description")]
    if proj_list:
        section_heading("Projects")
        for proj in proj_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r1 = p.add_run(proj.get("name", ""))
            set_font(r1, main_font, 11, bold=True)
            if proj.get("tech_stack"):
                r2 = p.add_run(f"   {proj['tech_stack']}")
                set_font(r2, main_font, 9.5, italic=True, color_hex="666666")
            body_para(proj.get("description", ""), indent=True)

    # Education
    edu_list = [e for e in data.get("education", []) if e.get("degree") or e.get("university")]
    if edu_list:
        section_heading("Education")
        for edu in edu_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r1 = p.add_run(edu.get("university", ""))
            set_font(r1, main_font, 11, bold=True)
            r2 = p.add_run(f"   {edu.get('year', '')}")
            set_font(r2, main_font, 9.5, italic=True, color_hex="666666")
            deg_para = doc.add_paragraph()
            dr = deg_para.add_run(edu.get("degree", ""))
            set_font(dr, main_font, 10, color_hex="333333")
            deg_para.paragraph_format.space_after = Pt(1)

    # Skills
    if data.get("skills", "").strip():
        section_heading("Skills")
        body_para(data["skills"])

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 6. Streamlit UI
# ==========================================
st.set_page_config(page_title="Resumed | AI Builder", layout="wide", page_icon="📄")

st.markdown("""
<style>
    .block-container { padding-top: 2rem; }
    .stTabs [data-baseweb="tab"] { font-size: 15px; font-weight: 600; }
    div[data-testid="stDownloadButton"] button {
        border-radius: 8px; font-weight: 600;
    }
</style>
""", unsafe_allow_html=True)

st.title("📄 Resumed — AI Resume Builder")
st.caption("Paste your raw background. AI writes it professionally. Download as Word or PDF.")

with st.sidebar:
    st.header("⚙️ Settings")
    template_choice = st.selectbox("Template Style", ["FAANG / Modern", "Classic / Serif"])
    st.markdown("---")
    st.markdown("### How it works")
    st.markdown("1. Paste your raw background info\n2. AI formats & enhances it\n3. Download as Word or PDF")
    st.markdown("---")
    st.markdown("**Tips for best results:**")
    st.markdown("- Include name, contact info\n- List any jobs/internships\n- Mention projects if any\n- Include your education")

style_key = "faang" if template_choice == "FAANG / Modern" else "xyz"

if "resume_data" not in st.session_state:
    st.session_state.resume_data = None

tab1, tab2 = st.tabs(["📝 1. Enter Details", "👁️ 2. Preview & Download"])

with tab1:
    st.markdown("### Paste Your Raw Background")
    st.caption("Don't worry about formatting. AI handles everything—grammar, bullet points, professional tone. Just dump the facts.")
    raw_text = st.text_area(
        "Your raw info:",
        height=260,
        placeholder="E.g.:\nMy name is Priya Sharma. Email: priya@gmail.com, Phone: +91-9876543210\n\nI studied Computer Science at IIT Delhi, graduated 2023.\n\nI interned at TCS for 6 months as a backend developer — worked on REST APIs using Node.js and MongoDB.\n\nI built a project called SmartNotes: a React + Firebase note-taking app with real-time sync.\n\nSkills: Python, JS, React, Node.js, SQL."
    )

    if st.button("✨ Generate Resume", use_container_width=True, type="primary"):
        if raw_text.strip():
            with st.spinner("AI is crafting your professional resume..."):
                result = extract_details_with_ai(raw_text)
                if result:
                    st.session_state.resume_data = result
                    st.success("✅ Resume generated! Switch to the **Preview & Download** tab.")
        else:
            st.warning("Please paste some text first.")

with tab2:
    if st.session_state.resume_data:
        data = st.session_state.resume_data
        final_html = render_template(data, style=style_key)

        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            docx_bytes = generate_docx(data, style=style_key)
            st.download_button(
                "📄 Download Word (.docx)",
                data=docx_bytes,
                file_name=f"{data.get('name', 'Resume').replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with col2:
            pdf_bytes = generate_pdf_reportlab(data, style=style_key)
            st.download_button(
                "📥 Download PDF",
                data=pdf_bytes,
                file_name=f"{data.get('name', 'Resume').replace(' ', '_')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )

        st.markdown("---")
        st.subheader("Live Preview")
        with st.container(border=True):
            st.components.v1.html(final_html, height=900, scrolling=True)

    else:
        st.info("👈 Go to the **Enter Details** tab, paste your background, and generate your resume first.")
